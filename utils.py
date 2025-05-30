import os
import requests
from io import BytesIO
from docx import Document as DocxDocument
import google.generativeai as genai # Para Gemini

# --- Constantes e Configurações ---
GROQ_API_BASE_URL = "https://api.groq.com/openai/v1"
CHATVOLT_API_BASE_URL = "https://api.chatvolt.ai/agents" # Verifique se é o endpoint correto

ALLOWED_TEXT_EXTENSIONS = ["txt", "pdf", "docx"]
ALLOWED_AUDIO_EXTENSIONS = ["mp3", "wav", "m4a", "ogg"]

# Template do prompt completo para Chatvolt (você forneceu)
CHATVOLT_FULL_PROMPT_TEMPLATE = """### INÍCIO DO PROMPT

Peticionador Jurídico Estruturado

Você é um agente jurídico especialista em Direito Brasileiro, com domínio técnico direito material e processual. Sua tarefa é redigir petições iniciais completas, com rigor jurídico, clareza argumentativa e linguagem forense formal. A redação deve obedecer à estrutura padrão de peças processuais, utilizando sempre doutrina relevante e dispositivos legais vigentes, tais como o Código Civil, Código de Defesa do Consumidor, Código de Processo Civil, Constituição Federal e resoluções da ANAC ou agências reguladoras, conforme o caso.

ESTRUTURA OBRIGATÓRIA DA PETIÇÃO INICIAL

ENDEREÇAMENTO
Inicie com o endereçamento apropriado ao juízo competente, como por exemplo:
AO JUIZADO ESPECIAL CÍVEL DA COMARCA DE [CIDADE/ESTADO]

QUALIFICAÇÃO DAS PARTES
Apresente a qualificação completa da parte autora (nome, nacionalidade, profissão, CPF, e endereço completo) e, logo após, da parte ré, seja pessoa física ou jurídica.

NOME DA AÇÃO (EM CAIXA ALTA)
Indique a natureza da ação, como:
AÇÃO DE INDENIZAÇÃO POR DANOS MORAIS E MATERIAIS

1. DOS FATOS
Descreva os fatos de forma objetiva, impessoal e cronológica, com riqueza de detalhes fornecidos pelo usuário. O texto deve permitir a perfeita compreensão da dinâmica do conflito jurídico, destacando:

Situação inicial e expectativa da parte autora;
O que ocorreu de modo inesperado ou lesivo;
Quais foram os efeitos imediatos e mediatos da conduta da parte ré;
Eventuais prejuízos financeiros, físicos ou emocionais sofridos;
Ausência de solução satisfatória pela parte ré.

2. DO DIREITO
Divida os fundamentos jurídicos em subtópicos numerados. Cada subtópico deve conter:

TÍTULO DESCRITIVO (em CAIXA ALTA);
Argumentação jurídica baseada em doutrina e lei

Exemplo de estrutura:

2.1 DA RESPONSABILIDADE CIVIL OBJETIVA DA COMPANHIA AÉREA
(falha na prestação de serviço, art. 14 do CDC, fortuito interno, obrigação de resultado)

2.2 DO DIREITO À INFORMAÇÃO E ASSISTÊNCIA AO CONSUMIDOR
(omissão, Resolução 400/2016 da ANAC, dever de assistência, art. 6º CDC)

2.3 DOS DANOS MATERIAIS
(comprovação de gastos adicionais, nota fiscal, art. 927 do CC)

2.4 DOS DANOS MORAIS – ATRASO SUPERIOR A QUATRO HORAS
(abalo psicológico, frustração, jurisprudência STJ, dano in re ipsa)

2.5 DA VIOLAÇÃO AO CONTRATO DE TRANSPORTE
(inadimplemento, art. 734 do CC, responsabilidade contratual)

Adapte a quantidade de tópicos conforme o caso concreto, podendo incluir outros como: Tutela Provisória, Desvio Produtivo, Cláusulas Abusivas, entre outros.

3. DOS PEDIDOS
Enumere com clareza os pedidos formulados ao juízo, como:

3.1 A citação da parte ré para apresentar defesa no prazo legal;

3.2 A condenação ao pagamento de indenização por danos materiais no valor de R$ XXX,XX;

3.3 A condenação ao pagamento de indenização por danos morais no valor de R$ XX.XXX,XX;

3.4 A produção de todas as provas em direito admitidas, incluindo documental, testemunhal e depoimento pessoal da parte ré;

3.5 A procedência total dos pedidos, com a condenação da ré nos ônus da sucumbência.

Incluir a frase: Optase pela audiência de conciliação.
Indicar valor da causa: Dáse à causa o valor de R$ XX.XXX,XX (valor por extenso), conforme o art. 292 do CPC.

ENCERRAMENTO
Termos em que,
Pede deferimento.

[CIDADE], [DIA] de [MÊS] de [ANO].

8. ASSINATURA
NOME DO ADVOGADO
ADVOGADO – OAB/UF Nº XXXXX

COMO FUNCIONA
O usuário fornecerá um relato com os fatos principais, ou documentos anexos. Com base nesses elementos, você, IA, deverá:

Organizar os fatos de forma clara e jurídica;
Identificar os direitos violados e os fundamentos legais cabíveis;
Escrever uma petição inicial completa, com linguagem técnico jurídica e estrutura formal, conforme o modelo acima.

Crie parágrafos de tamanho médio e claros. Separe o texto em parágrafos.

Se a resposta petição completa não couber em uma resposta, escreva [CONTINUA...] ao final da parte que foi escrita.

Importante: Escreva petições extensas. Os parágrafos devem ser médios, nem muito grandes nem muito pequenos. Cada tópico deve ser dividido em vários parágrafos e não em um parágrafo só. A argumentação deve ser longa e citar artigos da lei brasileira. Não utilize formatação em negrito.

---
REGRAS PARA JURISPRUDÊNCIA:

Faça uma citação literal de uma jurisprudência para cada um dos três tópicos. Não escreva "Fonte". Apenas salte uma linha após a citação e coloque logo em seguida, na mesma formatação, a informação sobre o tribunal e julgado. Exemplo: (REsp n. 1.733.136/RO, relator Ministro Paulo de Tarso Sanseverino, Terceira Turma, julgado em 21/9/2021, DJe de 24/9/2021.) Sempre insira a ementa no início da citação da jurisprudência. A ementa consiste em um pequeno texto em caixa alta com um resumo do julgado.

Exemplo de estrutura e formatação a serem seguidos:

RECURSO ESPECIAL. DIREITO CIVIL, CONSUMIDOR E PROCESSUAL CIVIL. RESPONSABILIDADE CIVIL PELO FATO DO SERVIÇO. ATRASO DE VOO. PASSAGEIRO MENOR (15 ANOS). JULGAMENTO ESTENDIDO. 1. PROCESSUAL CIVIL. JULGAMENTO ESTENDIDO. REALIZAÇÃO DA EXTENSÃO DO JULGAMENTO NA MESMA SESSÃO EM QUE LEVADO O VOTO VISTA VENCIDO. INTERPRETAÇÃO DA LOCUÇÃO "SENDO POSSÍVEL" CONSTANTE NO ENUNCIADO DO §1º DO ART. 842 DO CPC. NECESSIDADE DE SALVAGUARDA DO DEVIDO PROCESSO LEGAL E DA AMPLA DEFESA.
1.1. Esta Corte Superior é chamada a dizer da correta interpretação da locução "sendo possível" constante no início do §1º do art. 942 do CPC, dispositivo a condicionar a realização do julgamento estendido na mesma sessão em que verificada a não unanimidade, e, ainda, acerca do direito à indenização pelo atraso de voo doméstico.
1.2. O legislador de 2015 estava imbuído do espírito que se fez evidenciar em multifárias passagens do CPC no sentido do primado do devido processo legal, e centrado, notadamente, no constitucional direito ao contraditório e à ampla defesa.
1.3. A regra do §1º do art. 942 do CPC é clara e expressa acerca da possibilidade de o julgamento estendido ocorrer na mesma sessão quando: a) os demais integrantes do colegiado, embora não tendo participado do julgamento anterior, estiveram presentes à sustentação oral, dando-se por habilitados para o julgamento estendido, ou, b) quando se possibilite ao advogado, agora em face da extensão do julgamento e inclusão de novos integrantes, a realização de sustentação oral.
1.4. Caso concreto em que não se possibilitou ao advogado do demandante, ora recorrente, sustentar oralmente, o que, assim, faria nulo o julgamento realizado.
1.5. Nulidade, porém, que pode ser superada ante a possibilidade de, no mérito, ser provido o recurso especial, alcançando-lhe o direito à indenização pretendida. 2. CONSUMIDOR. RESPONSABILIDADE CIVIL PELO FATO DO SERVIÇO. DEFEITO NA PRESTAÇÃO DE SERVIÇO DE TRANSPORTE AEREO. ATRASO DE VOO. SUBMISSÃO DE MENOR DESACOMPANHADO A AGUARDAR POR NOVE HORAS EM CIDADE DESCONHECIDA PELO EMBARQUE. ATERRISAGEM EM CIDADE DIVERSA DA ORIGINALMENTE CONTRATADA (100 KM DISTANTE). ANGÚSTIA A QUE OS PAIS E O MENOR FORAM SUBMETIDOS A CONFIGURAR O DANO MORAL.
2.1. Grave defeito na prestação de serviço de transporte aéreo com a entrega de passageiro menor (15 anos) não na cidade de destino previamente contratada (Cacoal/RO), mas em uma cidade desconhecida situada a 100 km do local de destino, onde seu pai estaria em horário por deveras avançado: 23:15h (Ji-Paraná/RO).
2.2. Incomensurável a situação de angústia e aflição imposta aos pais do passageiro menor por esse grave e flagrante defeito na prestação do serviço de transporte aéreo .
2.3. Sequer o fato de se ter alegadamente ofertado transporte ao menor entre as cidades de Ji-Paraná/RO para Cacoal/RO, sobreleva, pois é claro que o pai não confiaria mais na empresa que tanto já havia demonstrado descumprir com as suas obrigações, deixando o seu filho a espera de transporte por quase metade de um dia e, no último trecho, submetendo-o, durante a madrugada, a transporte por uma van para levá-lo para a cidade de destino, com um motorista desconhecido, não se sabe se com outros passageiros ou não, nas nada seguras rodovias brasileiras.
2.4. Não se tem dúvidas que o direito brasileiro experimentou um período de banalização da reparação dos danos morais, reconhecendo-se o direito a toda sorte de situações, muitas delas em que efetivamente não se estava a lidar com violações a interesses ligados à esfera da dignidade humana e/ou dos direitos de personalidade.
2.5. Não se pode descurar, no entanto, que, presentes os elementos a evidenciar mais do que mero aborrecimento em ficar em um hotel, alimentado, no aguardo de um voo, mas a angústia de um menor e dos seus pais a espera de um voo por mais de nove horas, e a sua submissão a se deslocar para cidade a 100km da cidade de destino para buscar o seu filho, é devida a indenização pelos danos morais e materiais.
3. RECURSO ESPECIAL PROVIDO.
(REsp n. 1.733.136/RO, relator Ministro Paulo de Tarso Sanseverino, Terceira Turma, julgado em 21/9/2021, DJe de 24/9/2021.)

### FIM DO PROMPT
"""

# --- Funções Utilitárias Gerais ---
def allowed_file(filename, allowed_extensions):
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in allowed_extensions

def create_docx_from_text(text_content, title="Petição Gerada"):
    document = DocxDocument()
    document.add_heading(title, level=1)
    for line in text_content.split('\n'):
        document.add_paragraph(line)
    bio = BytesIO()
    document.save(bio)
    bio.seek(0)
    return bio

def simulated_petition_generation(tipo_peticao, assunto_principal, partes, fatos, outras_info, doc_files_info=None, audio_file_info=None):
    # (Sua lógica de simulação existente, pode ser simplificada ou mantida)
    from datetime import datetime
    data_atual = datetime.now().strftime('%d de %B de %Y')
    return f"""
EXCELENTÍSSIMO(A) SENHOR(A) DOUTOR(A) JUIZ(A) DE DIREITO DA [VARA SIMULADA] DA COMARCA DE [COMARCA SIMULADA].

Partes: {partes}
Tipo: {tipo_peticao.upper()}
Assunto: {assunto_principal.upper()}

1. DOS FATOS (SIMULADO)
{fatos}

2. DO DIREITO (SIMULADO)
[Fundamentação jurídica simulada...]
{outras_info}

3. DOS PEDIDOS (SIMULADO)
[Pedidos simulados...]

[CIDADE], {data_atual}.
[ADVOGADO SIMULADO]
OAB/UF [XXXXX]
    """.strip()

# --- Funções de API ---

# Groq
def build_groq_prompt(user_data):
    # Constrói um prompt mais detalhado para Groq baseado nos dados do usuário
    # (Similar ao que você já tinha, pode refinar mais)
    prompt = f"Você é um assistente jurídico especializado em criar rascunhos de petições no Brasil.\n"
    prompt += f"Tarefa: Gerar um rascunho de uma '{user_data['tipo_peticao']}' sobre '{user_data['assunto_principal']}'.\n"
    prompt += f"Partes envolvidas: {user_data['partes_str']}.\n"
    prompt += f"Descrição detalhada dos fatos: {user_data['fatos_str']}.\n"
    if user_data['outras_info_str']:
        prompt += f"Diretrizes adicionais ou pedidos específicos: {user_data['outras_info_str']}.\n"
    if user_data['documentos_texto']:
        prompt += "Conteúdo de documentos anexos (simulado):\n"
        for doc in user_data['documentos_texto']:
            prompt += f"- {doc['filename']}: {doc['content']}\n"
    prompt += "\nPor favor, gere o rascunho da petição solicitado, estruturando-o adequadamente com seções como 'DOS FATOS', 'DOS FUNDAMENTOS JURÍDICOS', 'DOS PEDIDOS', etc. Adapte o tom e a formalidade ao tipo de peça jurídica. Use linguagem forense formal e cite dispositivos legais brasileiros."
    return prompt

def query_groq_api(api_key, model_id, messages_history, temperature=0.7, max_tokens=3500):
    # (Sua função query_groq_api existente)
    if not api_key or not model_id: return "Erro: Chave API Groq ou modelo não configurado."
    url = f"{GROQ_API_BASE_URL}/chat/completions"
    headers = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}
    payload = {"model": model_id, "messages": messages_history, "temperature": temperature, "max_tokens": max_tokens}
    try:
        response = requests.post(url, headers=headers, json=payload, timeout=120)
        response.raise_for_status()
        json_response = response.json()
        if json_response.get("choices") and len(json_response["choices"]) > 0:
            content = json_response["choices"][0].get("message", {}).get("content")
            return content.strip() if content else "Erro: Resposta Groq sem conteúdo."
        return "Erro: Resposta Groq em formato inesperado."
    except requests.exceptions.HTTPError as http_err:
        return f"Erro HTTP da API Groq: {http_err} - {response.text if 'response' in locals() else ''}"
    except Exception as e: return f"Erro na API Groq: {e}"

# Chatvolt
def query_chatvolt_agent_with_template(api_key, agent_id, user_query_data, prompt_template):
    """
    Envia uma consulta para um agente Chatvolt usando um template de prompt.
    A API do Chatvolt pode ter uma maneira específica de lidar com templates.
    Esta é uma suposição de como poderia funcionar: o agente no Chatvolt
    foi configurado para entender este template e preencher as variáveis.
    Ou, formatamos o prompt aqui antes de enviar.

    Para este exemplo, vamos assumir que o 'user_query_data' são os fatos
    e o 'prompt_template' é o seu prompt extenso que já contém placeholders
    que o agente Chatvolt (ou um pré-processamento aqui) preencheria.

    Se o Chatvolt não suportar a passagem direta de um template e dados variáveis,
    você precisaria formatar o prompt_template com user_query_data *antes* de chamar esta função,
    e então passar o prompt finalizado como 'query'.

    Alternativa: O Chatvolt pode ter campos para "system_prompt" e "user_query".
    Nesse caso, prompt_template seria o system_prompt.
    """
    if not api_key or not agent_id:
        return "Erro: Chave API Chatvolt ou ID do Agente não configurados."

    url = f"{CHATVOLT_API_BASE_URL}/{agent_id}/query" # Verifique a URL correta da API Chatvolt
    headers = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}
    
    # O Chatvolt pode esperar um prompt completo ou dados para preencher um template no lado deles.
    # Aqui, vamos simular que o prompt_template é a instrução principal
    # e user_query_data são os fatos/contexto específico do caso.
    # O agente Chatvolt deve ser configurado para entender essa estrutura.
    # A forma mais simples para o Chatvolt é se o prompt_template já contém placeholders
    # e a API do Chatvolt permite enviar variáveis para preenchê-los, ou se o agente
    # é inteligente o suficiente para extrair os dados do user_query_data e usá-los
    # dentro do contexto do prompt_template.

    # Para usar seu prompt extenso como base:
    # O mais provável é que o seu "prompt_template" seja o "system prompt" ou instrução principal
    # e "user_query_data" seja a entrada do usuário (os fatos).
    # A query final enviada para o Chatvolt seria algo como:
    # query = f"{prompt_template}\n\nDados do caso fornecidos pelo usuário:\n{user_query_data}"
    # No entanto, a API do Chatvolt tem um campo "query". Vamos assumir que o agente
    # no Chatvolt já tem o "Peticionador Jurídico Estruturado" como seu system prompt.
    # Então, a "query" seria os dados do caso.

    data = {
        "query": user_query_data, # Dados específicos do caso
        "streaming": False,
        # Se o Chatvolt suportar, você pode enviar o prompt_template como um parâmetro adicional
        # ou garantir que o agente no Chatvolt já está configurado com ele.
        # "system_prompt": prompt_template (se a API suportar)
    }
    # Se o seu agente Chatvolt está configurado com o prompt extenso como System Prompt,
    # então a "query" deve ser apenas os dados variáveis (fatos, tipo de petição, etc.)
    # Se não, você precisa combinar o prompt_template e user_query_data na "query".
    # Para este exemplo, vou assumir que o prompt_template (seu prompt extenso)
    # precisa ser combinado com user_query_data.
    
    # O agente Chatvolt foi configurado COM o prompt que você forneceu.
    # Então, a 'query' deve ser os dados que o prompt espera que o usuário forneça.
    final_query_for_chatvolt = user_query_data # Os fatos e detalhes do caso

    data_chatvolt = {
        "query": final_query_for_chatvolt,
        "streaming": False
        # Se houver um campo específico para o prompt de sistema no Chatvolt,
        # você usaria o prompt_template lá. Se não, o prompt_template deve estar
        # configurado diretamente no agente na plataforma Chatvolt.
    }

    try:
        response = requests.post(url, headers=headers, json=data_chatvolt, timeout=180)
        response.raise_for_status()
        # A resposta do Chatvolt pode variar. Ajuste conforme a documentação.
        # Exemplo: response.json() pode ser {"response": "texto da petição", "conversationId": "..."}
        return response.json() # Espera-se que contenha uma chave 'response' com o texto.
    except requests.exceptions.HTTPError as http_err:
        return f"Erro HTTP da API Chatvolt: {http_err} - {response.text if 'response' in locals() else ''}"
    except Exception as e:
        return f"Erro na API Chatvolt: {str(e)}"


# Gemini
def query_gemini_api(api_key, model_name, prompt_text, max_output_tokens=8000):
    if not api_key: return "Erro: Chave da API Gemini não fornecida."
    try:
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel(
            model_name,
            generation_config=genai.types.GenerationConfig(
                max_output_tokens=max_output_tokens,
                temperature=0.7 # Ajuste conforme necessário
            ),
            # safety_settings ajustados para serem menos restritivos, CUIDADO em produção.
            safety_settings=[
                {"category": "HARM_CATEGORY_HARASSMENT", "threshold": "BLOCK_NONE"},
                {"category": "HARM_CATEGORY_HATE_SPEECH", "threshold": "BLOCK_NONE"},
                {"category": "HARM_CATEGORY_SEXUALLY_EXPLICIT", "threshold": "BLOCK_NONE"},
                {"category": "HARM_CATEGORY_DANGEROUS_CONTENT", "threshold": "BLOCK_NONE"},
            ]
        )
        response = model.generate_content(prompt_text)
        if response.parts:
            return response.text
        else:
            # Se não há 'parts', pode ser um bloqueio ou outro problema.
            # Imprime o motivo do bloqueio, se houver.
            print(f"Resposta Gemini sem 'parts'. Bloqueio? {response.prompt_feedback}")
            if response.prompt_feedback and response.prompt_feedback.block_reason:
                return f"Erro: Conteúdo bloqueado pela API Gemini. Motivo: {response.prompt_feedback.block_reason_message or response.prompt_feedback.block_reason}"
            return "Erro: Resposta da API Gemini vazia ou em formato inesperado."

    except Exception as e:
        error_message = f"Erro ao contatar a API Gemini: {str(e)}"
        if "API key not valid" in str(e):
            error_message = "Erro: Chave da API Gemini inválida ou não configurada corretamente."
        elif "quota" in str(e).lower():
            error_message = "Erro: Cota da API Gemini excedida. Tente novamente mais tarde."
        elif "resource_exhausted" in str(e).lower():
             error_message = "Erro: Recursos da API Gemini esgotados (provavelmente cota). Tente novamente mais tarde."
        print(error_message) # Log do erro
        return error_message

# --- Prompts para o Fluxo Gemini ---

# Instrução base para todos os prompts Gemini (persona e regras gerais)
GEMINI_BASE_INSTRUCTION = """
Você é um assistente jurídico especialista em Direito Brasileiro, com domínio técnico de direito material e processual.
Sua tarefa é auxiliar na redação de partes de uma petição inicial completa, com rigor jurídico, clareza argumentativa e linguagem forense formal.
Siga estritamente as instruções para cada parte da petição.
Não utilize formatação em negrito no texto da petição.
Crie parágrafos de tamanho médio e claros. Separe o texto em parágrafos.
A argumentação deve ser longa e citar artigos da lei brasileira quando pertinente.
"""

GEMINI_JURISPRUDENCE_INSTRUCTION = """
REGRAS PARA JURISPRUDÊNCIA (quando solicitado especificamente para um tópico de direito):
Faça uma citação literal de UMA jurisprudência relevante para o tópico.
Sempre insira a EMENTA EM CAIXA ALTA no início da citação da jurisprudência.
Após a citação, salte uma linha e coloque, na mesma formatação, a informação sobre o tribunal e julgado. Exemplo:
(REsp n. 1.733.136/RO, relator Ministro Paulo de Tarso Sanseverino, Terceira Turma, julgado em 21/9/2021, DJe de 24/9/2021.)
Não escreva "Fonte:" ou "Jurisprudência Citada:". Apenas a citação e os dados do julgado.
"""

def gemini_prompt_plan(user_data):
    return f"""{GEMINI_BASE_INSTRUCTION}
Considerando os seguintes dados fornecidos pelo usuário para uma petição inicial:
- Tipo de Peça Jurídica: {user_data['tipo_peticao']}
- Assunto Principal: {user_data['assunto_principal']}
- Partes Envolvidas: {user_data['partes_str']}
- Descrição dos Fatos: {user_data['fatos_str']}
- Outras Informações/Diretrizes: {user_data['outras_info_str']}

Tarefa: Crie um plano para a seção "2. DO DIREITO" da petição.
O plano deve consistir em uma lista numerada de TÍTULOS DESCRITIVOS (em CAIXA ALTA) para os subtópicos dos fundamentos jurídicos.
Cada título deve ser conciso e indicar o tema do respectivo subtópico.
Por exemplo:
2.1 DA RESPONSABILIDADE CIVIL OBJETIVA
2.2 DOS DANOS MORAIS
2.3 DA TUTELA DE URGÊNCIA

Não escreva o conteúdo dos tópicos, apenas a lista de títulos planejados para a seção "2. DO DIREITO".
"""

def gemini_prompt_addressing_facts(user_data):
    return f"""{GEMINI_BASE_INSTRUCTION}
Dados para a petição:
- Tipo de Peça Jurídica: {user_data['tipo_peticao']}
- Assunto Principal: {user_data['assunto_principal']}
- Partes Envolvidas (Autor, Réu, etc.): {user_data['partes_str']} (Use esta informação para a qualificação)
- Descrição Detalhada dos Fatos: {user_data['fatos_str']}
- Outras Informações/Diretrizes: {user_data['outras_info_str']}

Tarefa: Redija as seguintes seções da petição inicial:
1. ENDEREÇAMENTO (Ex: AO JUIZADO ESPECIAL CÍVEL DA COMARCA DE [CIDADE/ESTADO])
2. QUALIFICAÇÃO DAS PARTES (Apresente a qualificação completa da parte autora e da parte ré, com base nas informações fornecidas em "Partes Envolvidas". Se os detalhes não forem completos, use placeholders como [Nacionalidade], [Profissão], [CPF], [Endereço Completo], [CNPJ se aplicável], etc.)
3. NOME DA AÇÃO (EM CAIXA ALTA. Ex: AÇÃO DE INDENIZAÇÃO POR DANOS MORAIS E MATERIAIS)
4. Seção "1. DOS FATOS" (Descreva os fatos de forma objetiva, impessoal e cronológica, com base na "Descrição Detalhada dos Fatos" fornecida. Permita a perfeita compreensão da dinâmica do conflito.)

Siga rigorosamente a estrutura e formatação indicadas.
"""

def gemini_prompt_law_topic(user_data, topic_title):
    return f"""{GEMINI_BASE_INSTRUCTION}
{GEMINI_JURISPRUDENCE_INSTRUCTION}

Dados do caso:
- Tipo de Peça Jurídica: {user_data['tipo_peticao']}
- Assunto Principal: {user_data['assunto_principal']}
- Partes Envolvidas: {user_data['partes_str']}
- Descrição dos Fatos: {user_data['fatos_str']}
- Outras Informações/Diretrizes: {user_data['outras_info_str']}

Tarefa: Desenvolva o conteúdo argumentativo para o seguinte tópico da seção "2. DO DIREITO":
Título do Tópico: {topic_title}

Instruções para este tópico:
- Apresente argumentação jurídica robusta, baseada em doutrina relevante e dispositivos legais vigentes (ex: Código Civil, CDC, CPC, CF, Resoluções de agências, etc.).
- Se este tópico permitir e for relevante (ex: dano moral, responsabilidade civil, etc.), inclua UMA citação de jurisprudência conforme as regras de formatação especificadas acima.
- Se o título do tópico for, por exemplo, "2.3 DOS DANOS MATERIAIS", foque em comprovar os gastos, citar artigos como o 927 do CC, etc.
- Se for "2.4 DOS DANOS MORAIS", aborde o abalo psicológico, frustração, etc.
- Adapte a argumentação especificamente para o "{topic_title}".
"""

def gemini_prompt_requests_closing(user_data, developed_law_topics_summary):
    return f"""{GEMINI_BASE_INSTRUCTION}
Dados do caso:
- Tipo de Peça Jurídica: {user_data['tipo_peticao']}
- Assunto Principal: {user_data['assunto_principal']}
- Partes Envolvidas: {user_data['partes_str']}
- Descrição dos Fatos: {user_data['fatos_str']}
- Outras Informações/Diretrizes: {user_data['outras_info_str']}
- Resumo dos tópicos de direito já desenvolvidos: {developed_law_topics_summary}

Tarefa: Redija as seguintes seções finais da petição inicial:
1. Seção "3. DOS PEDIDOS" (Enumere com clareza os pedidos formulados ao juízo. Ex: citação, condenação em danos materiais R$ XXX,XX, danos morais R$ XX.XXX,XX, produção de provas, procedência total, condenação em sucumbência).
2. Inclua a frase: "Opta-se pela realização de audiência de conciliação." (ou "Opta-se pela não realização de audiência de conciliação." se mais apropriado ou se indicado nas diretrizes).
3. Indique o VALOR DA CAUSA: "Dá-se à causa o valor de R$ [CALCULAR OU INDICAR PLACEHOLDER COM BASE NOS PEDIDOS E FATOS], conforme o art. 292 do CPC." (Se possível, sugira um valor ou um placeholder claro).
4. ENCERRAMENTO (Termos em que, Pede deferimento. [LOCAL], [DATA].)
5. ASSINATURA (NOME DO ADVOGADO, ADVOGADO – OAB/UF Nº XXXXX) (Use placeholders para nome, local, data e OAB).
"""

# --- Fluxo de Geração com Gemini ---
def generate_petition_gemini_flow(api_key, user_data, model_name="gemini-2.0-flash"):
    """
    Orquestra o fluxo de múltiplas chamadas à API Gemini para gerar a petição.
    """
    full_petition_parts = []
    max_tokens_per_step = 7500 # Ajuste conforme necessário e limites do modelo flash

    # Etapa 1: Planejamento dos tópicos de Direito
    prompt_plan = gemini_prompt_plan(user_data)
    plan_text = query_gemini_api(api_key, model_name, prompt_plan, max_output_tokens=1000)
    if plan_text.startswith("Erro"): return f"Erro no planejamento: {plan_text}"
    
    # Extrair os títulos do plano (isso pode precisar de um parseamento mais robusto)
    law_topics_titles = [line.strip() for line in plan_text.split('\n') if line.strip() and (line.strip().startswith("2.") or line.strip()[0].isdigit() and '.' in line)]
    if not law_topics_titles:
        # Fallback se o plano não veio como esperado, ou se a IA não listou tópicos
        # Poderia usar uma lista padrão ou pedir ao usuário, mas por ora, erro.
         return f"Erro: Não foi possível extrair os tópicos de direito do plano gerado pela IA. Plano recebido:\n{plan_text}"
    print(f"Plano de Tópicos do Direito (Gemini): {law_topics_titles}") # Log

    # Etapa 2: Endereçamento e Fatos
    prompt_address_facts = gemini_prompt_addressing_facts(user_data)
    address_facts_text = query_gemini_api(api_key, model_name, prompt_address_facts, max_output_tokens=max_tokens_per_step)
    if address_facts_text.startswith("Erro"): return f"Erro no endereçamento/fatos: {address_facts_text}"
    full_petition_parts.append(address_facts_text)

    # Etapa 3: Desenvolvimento de cada Tópico do Direito
    developed_law_sections_text = ["\n\n2. DO DIREITO\n"]
    for topic_title in law_topics_titles:
        if not topic_title: continue # Pular linhas vazias se houver
        prompt_law_topic = gemini_prompt_law_topic(user_data, topic_title)
        law_topic_text = query_gemini_api(api_key, model_name, prompt_law_topic, max_output_tokens=max_tokens_per_step)
        if law_topic_text.startswith("Erro"):
            developed_law_sections_text.append(f"\n--- ERRO AO GERAR TÓPICO: {topic_title} ---\n{law_topic_text}\n--- FIM DO ERRO ---\n")
            # Decide se continua ou para. Por ora, continua com a marcação do erro.
        else:
            developed_law_sections_text.append(f"\n{law_topic_text}\n") # Adiciona o título e o conteúdo do tópico
    
    full_petition_parts.append("".join(developed_law_sections_text))

    # Etapa 4: Pedidos e Encerramento
    # Criar um resumo dos tópicos de direito para o contexto dos pedidos
    developed_law_topics_summary = "; ".join(law_topics_titles)
    prompt_requests_closing = gemini_prompt_requests_closing(user_data, developed_law_topics_summary)
    requests_closing_text = query_gemini_api(api_key, model_name, prompt_requests_closing, max_output_tokens=max_tokens_per_step)
    if requests_closing_text.startswith("Erro"): return f"Erro nos pedidos/encerramento: {requests_closing_text}" # Ou anexa o erro
    full_petition_parts.append(f"\n{requests_closing_text}")

    return "\n\n".join(full_petition_parts)